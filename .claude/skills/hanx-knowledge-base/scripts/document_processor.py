#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Document Processor for Knowledge Base

Extracts text and metadata from various document formats:
- PDF (.pdf)
- Word Documents (.docx, .doc)
- Markdown (.md, .markdown)
- Plain Text (.txt)
- CSV (.csv)
- HTML (.html, .htm)
- JSON (.json)

Usage:
    from document_processor import DocumentProcessor, process_document

    # Process single document
    content, metadata = process_document("document.pdf")

    # Process with additional metadata
    content, metadata = process_document(
        "document.pdf",
        base_metadata={"category": "technical", "author": "John Doe"}
    )
"""

import os
import sys
from typing import Dict, Any, Tuple, List
from pathlib import Path
from datetime import datetime
from abc import ABC, abstractmethod

# Try importing dependencies with auto-install fallback
try:
    import PyPDF2
    import docx
    import csv
    import yaml
    import markdown
    from bs4 import BeautifulSoup
except ImportError:
    import subprocess
    print("[INSTALL] Installing document processing dependencies...")
    subprocess.check_call([
        sys.executable, "-m", "pip", "install",
        "PyPDF2", "python-docx", "pyyaml", "markdown",
        "beautifulsoup4", "lxml"
    ])
    import PyPDF2
    import docx
    import csv
    import yaml
    import markdown
    from bs4 import BeautifulSoup


class BaseProcessor(ABC):
    """Base class for document processors"""

    def __init__(self, base_metadata: Dict[str, Any] = None):
        """
        Initialize processor

        Args:
            base_metadata: Base metadata to include for all processed documents
        """
        self.base_metadata = base_metadata or {}

    @abstractmethod
    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process document and extract content + metadata

        Args:
            file_path: Path to document file

        Returns:
            Tuple of (content, metadata)
        """
        pass

    def _get_base_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Get base metadata for a file

        Args:
            file_path: Path to file

        Returns:
            Dictionary with file metadata
        """
        path = Path(file_path)
        metadata = self.base_metadata.copy()

        stat = path.stat()
        metadata.update({
            "source": str(path.absolute()),
            "file_name": path.name,
            "file_type": path.suffix.lower()[1:] if path.suffix else "unknown",
            "file_size_bytes": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        })

        return metadata


class PDFProcessor(BaseProcessor):
    """Process PDF documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from PDF

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)

            # Extract text from all pages
            content_parts = []
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text.strip():
                    content_parts.append(f"[Page {page_num}]\n{text}")

            content = "\n\n".join(content_parts)

            # Extract PDF metadata
            if pdf.metadata:
                pdf_meta = {}
                for key, value in pdf.metadata.items():
                    # Remove leading slash from keys
                    clean_key = key.lstrip('/').lower()
                    pdf_meta[clean_key] = str(value) if value else ""

                metadata.update(pdf_meta)

            # Add page count
            metadata['page_count'] = len(pdf.pages)

        return content.strip(), metadata


class DocxProcessor(BaseProcessor):
    """Process Word documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        doc = docx.Document(file_path)

        # Extract core properties
        core_props = doc.core_properties
        if core_props:
            prop_dict = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
            }

            # Add dates if available
            if core_props.created:
                prop_dict["doc_created"] = core_props.created.isoformat()
            if core_props.modified:
                prop_dict["doc_modified"] = core_props.modified.isoformat()

            # Filter out empty values
            metadata.update({k: v for k, v in prop_dict.items() if v})

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # Extract text from tables
        if doc.tables:
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_texts.append(row_text)

            if table_texts:
                content += "\n\n[Tables]\n" + "\n".join(table_texts)

        metadata['paragraph_count'] = len(paragraphs)
        metadata['table_count'] = len(doc.tables)

        return content.strip(), metadata


class MarkdownProcessor(BaseProcessor):
    """Process Markdown documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from Markdown

        Args:
            file_path: Path to Markdown file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

            # Check for YAML frontmatter
            if content.startswith('---'):
                end_idx = content.find('---', 3)
                if end_idx != -1:
                    frontmatter = content[3:end_idx]
                    try:
                        yaml_metadata = yaml.safe_load(frontmatter)
                        if yaml_metadata and isinstance(yaml_metadata, dict):
                            metadata.update(yaml_metadata)
                        content = content[end_idx + 3:].strip()
                    except yaml.YAMLError:
                        pass  # Invalid YAML, skip frontmatter parsing

        return content.strip(), metadata


class TextProcessor(BaseProcessor):
    """Process plain text documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from TXT

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # Add basic stats
        lines = content.split('\n')
        metadata['line_count'] = len(lines)
        metadata['word_count'] = len(content.split())
        metadata['char_count'] = len(content)

        return content.strip(), metadata


class HTMLProcessor(BaseProcessor):
    """Process HTML documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from HTML

        Args:
            file_path: Path to HTML file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')

            # Extract metadata from meta tags
            for tag in soup.find_all('meta'):
                name = tag.get('name', '').lower()
                content = tag.get('content', '')
                if name and content:
                    metadata[name] = content

            # Extract title
            if soup.title:
                metadata['title'] = soup.title.string

            # Remove script, style, and other non-content tags
            for tag in soup(['script', 'style', 'meta', 'link', 'noscript']):
                tag.decompose()

            # Extract text
            content = soup.get_text(separator='\n', strip=True)

            # Clean up multiple newlines
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            content = '\n\n'.join(lines)

        return content.strip(), metadata


class CSVProcessor(BaseProcessor):
    """Process CSV documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from CSV

        Args:
            file_path: Path to CSV file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'r', encoding='utf-8') as file:
            # Detect delimiter
            sample = file.read(4096)
            file.seek(0)
            sniffer = csv.Sniffer()
            try:
                delimiter = sniffer.sniff(sample).delimiter
            except:
                delimiter = ','

            reader = csv.DictReader(file, delimiter=delimiter)
            headers = reader.fieldnames
            metadata['headers'] = list(headers) if headers else []

            # Convert rows to text
            content_lines = []
            row_count = 0

            for row in reader:
                row_count += 1
                # Format each row as "key: value | key: value ..."
                row_text = " | ".join([f"{k}: {v}" for k, v in row.items()])
                content_lines.append(row_text)

            metadata['row_count'] = row_count
            metadata['column_count'] = len(headers) if headers else 0

            # Add header as first line
            if headers:
                header_line = " | ".join(headers)
                content_lines.insert(0, f"[Headers]\n{header_line}\n")

            content = "\n".join(content_lines)

        return content.strip(), metadata


class JSONProcessor(BaseProcessor):
    """Process JSON documents"""

    def process(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from JSON

        Args:
            file_path: Path to JSON file

        Returns:
            Tuple of (content, metadata)
        """
        metadata = self._get_base_metadata(file_path)

        with open(file_path, 'r', encoding='utf-8') as file:
            data = yaml.safe_load(file)  # yaml.safe_load also handles JSON

            # Convert JSON to readable text
            content = yaml.dump(data, default_flow_style=False, sort_keys=False)

            # Add structure info to metadata
            if isinstance(data, dict):
                metadata['json_type'] = 'object'
                metadata['top_level_keys'] = list(data.keys())
            elif isinstance(data, list):
                metadata['json_type'] = 'array'
                metadata['array_length'] = len(data)
            else:
                metadata['json_type'] = 'primitive'

        return content.strip(), metadata


# Processor registry
PROCESSORS = {
    '.pdf': PDFProcessor,
    '.docx': DocxProcessor,
    '.doc': DocxProcessor,
    '.md': MarkdownProcessor,
    '.markdown': MarkdownProcessor,
    '.txt': TextProcessor,
    '.text': TextProcessor,
    '.html': HTMLProcessor,
    '.htm': HTMLProcessor,
    '.csv': CSVProcessor,
    '.json': JSONProcessor,
}


def get_processor(
    file_path: str,
    base_metadata: Dict[str, Any] = None
) -> BaseProcessor:
    """
    Get appropriate processor for file type

    Args:
        file_path: Path to file
        base_metadata: Optional base metadata

    Returns:
        Processor instance

    Raises:
        ValueError: If file type is not supported
    """
    ext = Path(file_path).suffix.lower()

    processor_class = PROCESSORS.get(ext)
    if processor_class is None:
        # Default to text processor for unknown types
        print(f"[WARNING] Unknown file type '{ext}', using text processor")
        processor_class = TextProcessor

    return processor_class(base_metadata)


def process_document(
    file_path: str,
    base_metadata: Dict[str, Any] = None
) -> Tuple[str, Dict[str, Any]]:
    """
    Process a document and extract content + metadata

    Args:
        file_path: Path to document
        base_metadata: Optional base metadata to include

    Returns:
        Tuple of (content, metadata)

    Example:
        content, metadata = process_document("document.pdf")
        print(f"Title: {metadata.get('title', 'N/A')}")
        print(f"Content: {content[:200]}...")
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_path.is_file():
        raise ValueError(f"Not a file: {file_path}")

    processor = get_processor(str(file_path), base_metadata)
    return processor.process(str(file_path))


def process_directory(
    directory_path: str,
    recursive: bool = True,
    base_metadata: Dict[str, Any] = None,
    file_extensions: List[str] = None
) -> List[Tuple[str, Dict[str, Any]]]:
    """
    Process all documents in a directory

    Args:
        directory_path: Path to directory
        recursive: Process subdirectories
        base_metadata: Base metadata for all documents
        file_extensions: Filter by file extensions (e.g., ['.pdf', '.docx'])

    Returns:
        List of (content, metadata) tuples
    """
    directory = Path(directory_path)

    if not directory.exists() or not directory.is_dir():
        raise ValueError(f"Invalid directory: {directory_path}")

    results = []
    pattern = "**/*" if recursive else "*"

    for file_path in directory.glob(pattern):
        if not file_path.is_file():
            continue

        # Filter by extension if specified
        if file_extensions:
            if file_path.suffix.lower() not in file_extensions:
                continue

        # Check if processor exists for this file type
        if file_path.suffix.lower() not in PROCESSORS:
            continue

        try:
            print(f"[PROCESSING] {file_path.name}...")
            content, metadata = process_document(str(file_path), base_metadata)
            results.append((content, metadata))
        except Exception as e:
            print(f"[ERROR] Failed to process {file_path.name}: {e}")

    return results


# Example usage
if __name__ == "__main__":
    print("=" * 80)
    print("Document Processor - Test")
    print("=" * 80)
    print()

    # Test with a text file
    test_file = "test_document.txt"

    # Create test file
    with open(test_file, 'w') as f:
        f.write("This is a test document.\n\nIt has multiple paragraphs.")

    try:
        print(f"[TEST] Processing: {test_file}")
        content, metadata = process_document(test_file)

        print("\n[METADATA]")
        for key, value in metadata.items():
            print(f"  {key}: {value}")

        print(f"\n[CONTENT]")
        print(content)

        print("\n✅ Test complete!")

    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
